import pandas as pd
import docx2txt
import os
from docx import Document
import glob
import csv
import shutil
import re
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
import zipfile
import tempfile



def rename_file(old_name, new_name):
    try:
        os.rename(old_name, new_name)
        print(f"File '{old_name}' renamed to '{new_name}' successfully.")
        return True
    except FileNotFoundError:
        print(f"File '{old_name}' not found.")
        return False
    except FileExistsError:
        print(f"File '{new_name}' already exists. Rename operation failed.")
        return False
    except Exception as e:
        print(f"An error occurred: {str(e)}")
        return False




def xlsx_to_csv(xlsx_file, csv_file, col_1, col_2, col_3):
    # convert from xlsx to csv
    exl = pd.read_excel(xlsx_file)
    
    exl.to_csv(csv_file, index=False)



    # Read the CSV file
    df = pd.read_csv(csv_file)

    # Extract data from the 'donor name' column
    df.columns = [col_1, col_2, col_3]

    # Remove rows with missing data in either column.
    df = df.dropna(subset=[col_1, col_2, col_3])

    name_data = df[col_1].tolist()

    # Extract data from the 'donor amount' column
    amount_data = df[col_2].tolist()

    return name_data, amount_data


def convert_doc_txt(docx_file, txt_file):
    MY_TEXT = docx2txt.process(docx_file)
    with open(txt_file, "w") as text_file:
        print(MY_TEXT, file=text_file)


def replace_words(name, amount):
    # Read the text file
    with open('template.txt', 'r') as file:
        content = file.read()

    # Replace the "donor name" and "donor amount" with your desired text
    content_modified = content.replace('Donor Name', name)
    content_modified = content_modified.replace('Donor Amount', amount)

    # Create a new file with a dynamic name (e.g., name.txt) and store the modified content there
    new_filename = f'{name}.txt'
    with open(new_filename, 'w') as new_file:
        new_file.write(content_modified)

    print(f'Text successfully replaced and saved to {new_filename}!')



def replace_content_after_date(docx_filename, txt_filename, output_filename):
    # Read the content from the text file
    with open(txt_filename, 'r', encoding='utf-8') as txt_file:
        new_content = txt_file.read()

    # Open the existing .docx file
    doc = Document(docx_filename)
    
    # Flag to indicate when to start replacing content
    start_replacing = False
    
    # Iterate through paragraphs and replace content after finding the date
    for paragraph in doc.paragraphs:
        if "27th September 2023" in paragraph.text:
            start_replacing = True
        
        # If the flag is set, clear the paragraph and add the new content
        if start_replacing:
            for run in paragraph.runs:
                run.clear()
            paragraph.add_run(new_content)
        break
    # Save the modified document with a new name
    doc.save(output_filename)

    print(f'Content after "27th September 2023" has been replaced in {output_filename}.')

def replace_words_in_docx(docx_filename, replacements):
    # Load the .docx document
    doc = Document(docx_filename)

    # Iterate through paragraphs
    for paragraph in doc.paragraphs:
        for old_word, new_word in replacements.items():
            # Replace the old word with the new word in the paragraph
            paragraph.text = paragraph.text.replace(old_word, new_word)

    # Save the modified document with a new name
    output_filename = f"{docx_filename}"
    doc.save(output_filename)

    print(f"Words replaced and saved in {output_filename}")

def read_csv(filename):
    names = []
    amounts = []

    with open(filename, mode='r', newline='', encoding='utf-8') as file:
        reader = csv.DictReader(file)
        
        # Check if the headers 'Donor Name' and 'Donor Amount' exist in the CSV file
        if 'Donor Name' in reader.fieldnames and 'Donor Amount' in reader.fieldnames:
            for row in reader:
                name = row['Donor Name'].strip()
                amount_str = row['Donor Amount'].strip()
                
                # Handle missing or empty amount values
                if amount_str:
                    amounts.append(amount_str)
                else:
                    amounts.append("0")
                
                names.append(name)
        else:
            print("CSV file does not contain the expected column headers.")

    return names, amounts

def move_docx_files_to_thank_you_folder(source_folder, destination_folder):
    # Ensure the "Thank you" folder exists; create it if it doesn't
    if not os.path.exists(destination_folder):
        os.makedirs(destination_folder)

    # Iterate through files in the source folder
    for filename in os.listdir(source_folder):
        source_path = os.path.join(source_folder, filename)

        # Check if the file is a .docx file
        if filename.endswith(".docx"):
            destination_path = os.path.join(destination_folder, filename)
            if filename == "template.docx":
                continue


            # Move the .docx file to the "Thank you" folder
            shutil.move(source_path, destination_path)
            print(f'Moved {filename} to {destination_folder}')
    
def clean_csv_data(input_csv_path, output_csv_path):
    # Read the CSV file into a pandas DataFrame, skipping empty lines and header rows.
    df = pd.read_csv(input_csv_path, skip_blank_lines=True, header=None, names=["Donor Name", "Donor Amount"])

    # Remove rows with missing Donor Name or Donor Amount.
    df = df.dropna(subset=["Donor Name", "Donor Amount"], how="any")

    # Remove rows with "Unnamed: 1" in the "Donor Name" column.
    df = df[df["Donor Name"] != "Unnamed: 1"]

    # Remove rows with "Unnamed: 2" in the "Donor Amount" column.
    df = df[df["Donor Amount"] != "Unnamed: 2"]

    # Remove duplicate entries in "Donor Name" and "Donor Amount" columns.
    # df = df.drop_duplicates(subset=["Donor Name", "Donor Amount"])/
    # df = df.drop_duplicates(subset=["Donor Name", "Donor Amount"], keep="first")
    # df = df[df.duplicated(subset=["Donor Name", "Donor Amount"]) == False]
    # df = df[~df.duplicated(subset=["Donor Name", "Donor Amount"], keep="first")]
    # df["Donor Name"] = df["Donor Name"].str.strip()
    # df["Donor Amount"] = df["Donor Amount"].str.strip()
    # df = df[~df.duplicated(subset=["Donor Name", "Donor Amount"], keep="first")]
    def clean_and_normalize(s):
        s = re.sub(r'[^a-zA-Z0-9\s]', '', s)  # Remove special characters
        s = s.strip()  # Trim leading/trailing whitespace
        return s

    # Clean and normalize "Donor Name" and "Donor Amount" columns.
    df["Donor Name"] = df["Donor Name"].apply(clean_and_normalize)
    df["Donor Amount"] = df["Donor Amount"].apply(clean_and_normalize)

    # Remove rows where both "Donor Name" and "Donor Amount" are duplicates.
    df = df[~df.duplicated(subset=["Donor Name", "Donor Amount"], keep="first")]


    # Save the cleaned data to a new CSV file.
    df.to_csv(output_csv_path, index=False)


def delete_all_txt_files(directory_path):
    try:
        # Ensure the directory path exists
        if not os.path.exists(directory_path):
            raise FileNotFoundError(f"Directory not found: {directory_path}")

        # List all TXT files in the directory
        txt_files = [file for file in os.listdir(directory_path) if file.endswith(".txt")]

        for txt_file in txt_files:
            # Construct the full file path
            file_path = os.path.join(directory_path, txt_file)

            # Delete the file
            os.remove(file_path)

            print(f"Deleted: {txt_file}")

        print(f"Deleted {len(txt_files)} TXT files.")
    except Exception as e:
        print(f"An error occurred: {e}")





app = FastAPI()

origins = ["*"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
@app.post("/upload/")
async def upload_files(docx_file: UploadFile = File(...), xlsx_file: UploadFile = File(...)):
    try:
        # Create a directory to store uploaded files (if it doesn't exist)
        # if not os.path.exists("uploads"):
        #     os.makedirs("uploads")

        # Save the uploaded DOCX file
        docx_file_path = os.path.join("./", "template.docx")
        with open(docx_file_path, "wb") as f:
            f.write(docx_file.file.read())

        # Save the uploaded XLSX file
        xlsx_file_path = os.path.join("./", "donations.xlsx")
        with open(xlsx_file_path, "wb") as f:
            f.write(xlsx_file.file.read())
        name_arr, amount_arr = xlsx_to_csv("donations.xlsx", "donations.csv","Donor Name", "Donor Amount", "")
        clean_csv_data("donations.csv", "donations.csv")
        convert_doc_txt("template.docx", "template.txt")
    



        filename = 'donations.csv'  # Replace with your CSV file's name
        names, amounts = read_csv(filename)

        # # # # Print the names and amounts
        for name, amount in zip(names, amounts):
                text_file = name+".txt"
                docx_folder = "completed/"+ name + "." + "docx"
                docx_file = name + ".docx"
                cleaned_amount = amount.replace('₦', '').replace(',', '')
        
                try:
                    amount = int(cleaned_amount)
                except ValueError:
                    # Handle invalid amount value here
                    print(f"Invalid amount value for {name}: {amount}")
                    continue

                # Format the amount with commas
                formatted_amount = '{:,.0f}'.format(amount)

                # formatted_amount = '{:,}'.format(int(amount))
                replacements = {
                    "Donor Name": name ,
                    "Donor Amount": formatted_amount}
                replace_words(name, str(amount))
                replace_content_after_date("template.docx", text_file, docx_file)
                replace_words_in_docx(docx_file, replacements)
        
        move_docx_files_to_thank_you_folder(".", "completed")
        delete_all_txt_files(".")

        return JSONResponse(content={
            "message": "Files uploaded successfully and modified",
            "download_link": "http://127.0.0.1:8000/download/completed"
            }, status_code=200)
    except Exception as e:
         return JSONResponse(content={"error": str(e)}, status_code=500)


@app.get("/download/{folder_name}")
async def download_folder(folder_name: str):
    try:
        folder_path = os.path.join("./", folder_name)

        if not os.path.exists(folder_path):
            raise HTTPException(status_code=404, detail="Folder not found")

        # Create a temporary directory to store the zip archive
        temp_dir = tempfile.mkdtemp()

        # Create a zip file to store the folder contents
        zip_filename = f"{folder_name}.zip"
        zip_filepath = os.path.join(temp_dir, zip_filename)

        with zipfile.ZipFile(zip_filepath, "w", zipfile.ZIP_DEFLATED) as zipf:
            for root, _, files in os.walk(folder_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, folder_path)
                    zipf.write(file_path, arcname=arcname)

        # Serve the zip archive for download
        return FileResponse(zip_filepath, media_type='application/zip', filename=zip_filename)
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)



# if __name__ == "__main__":
#     name_arr, amount_arr = xlsx_to_csv("Donations.xlsx", "donations.csv","Donor Name", "Donor Amount", "")
#     clean_csv_data("donations.csv", "donations.csv")
#     convert_doc_txt("template.docx", "template.txt")
    



#     filename = 'donations.csv'  # Replace with your CSV file's name
#     names, amounts = read_csv(filename)

#     # # # # Print the names and amounts
#     for name, amount in zip(names, amounts):
#             text_file = name+".txt"
#             docx_folder = "Thank You/"+ name + "." + "docx"
#             docx_file = name + ".docx"
#             cleaned_amount = amount.replace('₦', '').replace(',', '')
    
#             try:
#                 amount = int(cleaned_amount)
#             except ValueError:
#                 # Handle invalid amount value here
#                 print(f"Invalid amount value for {name}: {amount}")
#                 continue

#             # Format the amount with commas
#             formatted_amount = '{:,.0f}'.format(amount)

#             # formatted_amount = '{:,}'.format(int(amount))
#             replacements = {
#                 "Donor Name": name ,
#                 "Donor Amount": formatted_amount}
#             replace_words(name, str(amount))
#             replace_content_after_date("template.docx", text_file, docx_file)
#             replace_words_in_docx(docx_file, replacements)
    
#     move_docx_files_to_thank_you_folder(".", "Thank You")
#     delete_all_txt_files(".")
#     # replace_content_after_date("thank.docx", "Dejo Ajani.txt", "Dejo Ajani.docx")
# # if __name__ == "__main__":
# #     name_arr, amount_arr = xlsx_to_csv("Donations.xlsx", "donations.csv", "Donor Name", "Donor Amount")
# #     for name, amount in zip(name_arr, amount_arr):
# #         print("Name:", name)
# #         print("Amount:", amount)
    
    
#     # Specify filenames and output filename
#     # docx_filename = 'thank.docx'  # Replace with your input .docx filename
#     # txt_filename = 'your_file.txt'    # Replace with your input .txt filename
#     # output_filename = 'your.docx'  # Replace with the desired output filename

#     # # Call the function to replace content
#     # replace_content_after_date(docx_filename, txt_filename, output_filename)
